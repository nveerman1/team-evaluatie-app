from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Re-use colour palette & helpers from projectplan_export ───────────────────
from app.services.projectplan_export import (
    COLOR_DARK,
    COLOR_ACCENT,
    COLOR_MUTED,
    COLOR_LABEL,
    FONT_NAME,
    _BORDER_COLOR,
    _HEADER_FILL,
    _set_cell_shading,
    _set_cell_borders,
    _set_cell_margins,
    _add_paragraph_bottom_border,
    _spacing,
    _set_default_font,
)


# ── Internal helpers ──────────────────────────────────────────────────────────


def _add_rubric_header(
    doc: Document,
    assessment_title: str,
    team_number,
    team_members: list,
    export_date: str,
) -> None:
    """Top header: school label → doc type → assessment title → meta → hr."""
    school = doc.add_paragraph("TECHNASIUM MBH")
    school.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(school, before=0, after=2)
    run = school.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED
    run.font.all_caps = True

    doc_type = doc.add_paragraph("Beoordeling")
    _spacing(doc_type, before=4, after=0)
    run = doc_type.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK

    if assessment_title:
        proj = doc.add_paragraph(assessment_title)
        _spacing(proj, before=2, after=4)
        run = proj.runs[0]
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK

    parts: list[str] = []
    if team_number is not None:
        parts.append(f"Team {team_number}")
    if team_members:
        parts.append(", ".join(team_members))
    parts.append(f"Datum: {export_date}")
    meta = doc.add_paragraph("   ·   ".join(parts))
    _spacing(meta, before=0, after=2)
    run = meta.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED

    hr = doc.add_paragraph()
    _spacing(hr, before=4, after=8)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E293B")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_rubric_table(doc: Document, criteria: list, scores_map: dict) -> None:
    """Render the rubric table: grouped by category when categories are present."""
    has_categories = any(c.get("category") for c in criteria)

    # Build rows grouped by category
    if has_categories:
        from collections import OrderedDict

        cat_groups: dict[str, list] = OrderedDict()
        for c in criteria:
            cat = c.get("category") or "Overig"
            cat_groups.setdefault(cat, []).append(c)
    else:
        cat_groups = {"": criteria}

    # Create 3-column table
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"

    # Set approximate column widths via XML
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    # Header row
    hdr_row = table.add_row()
    for i, label in enumerate(("Criterium", "Score", "Toelichting")):
        cell = hdr_row.cells[i]
        cell.text = label
        _set_cell_shading(cell, _HEADER_FILL)
        _set_cell_borders(cell, _BORDER_COLOR)
        _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.name = FONT_NAME
            r.font.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_DARK

    _CATEGORY_FILL = "EFF6FF"  # blue-50

    for category, crit_list in cat_groups.items():
        # Category subheading row (only when categories are used)
        if has_categories and category:
            cat_row = table.add_row()
            merged = cat_row.cells[0].merge(cat_row.cells[1]).merge(cat_row.cells[2])
            merged.text = category
            _set_cell_shading(merged, _CATEGORY_FILL)
            _set_cell_borders(merged, _BORDER_COLOR)
            _set_cell_margins(merged, top=60, bottom=60, left=100, right=100)
            if merged.paragraphs[0].runs:
                r = merged.paragraphs[0].runs[0]
                r.font.name = FONT_NAME
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = COLOR_ACCENT

        for criterion in crit_list:
            crit_id = criterion["id"]
            score_obj = scores_map.get(crit_id, {})
            score_val = score_obj.get("score")
            comment_val = score_obj.get("comment") or ""

            data_row = table.add_row()
            values = [
                criterion["name"],
                str(score_val) if score_val is not None else "—",
                comment_val,
            ]
            for i, val in enumerate(values):
                cell = data_row.cells[i]
                cell.text = val
                _set_cell_borders(cell, _BORDER_COLOR)
                _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                if cell.paragraphs[0].runs:
                    r = cell.paragraphs[0].runs[0]
                    r.font.name = FONT_NAME
                    r.font.size = Pt(8)
                    r.font.color.rgb = COLOR_DARK

    spacer = doc.add_paragraph()
    _spacing(spacer, before=0, after=4)


def _add_score_summary(
    doc: Document,
    total_score,
    grade,
    general_comment: str | None,
) -> None:
    """Total score, grade, and optional general comment below the rubric table."""
    if total_score is not None or grade is not None:
        summary_parts = []
        if total_score is not None:
            summary_parts.append(f"Totaalscore: {total_score}")
        if grade is not None:
            summary_parts.append(f"Cijfer: {grade}")
        summary_para = doc.add_paragraph("   ·   ".join(summary_parts))
        _spacing(summary_para, before=2, after=4)
        for run in summary_para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = COLOR_DARK

    if general_comment:
        label_para = doc.add_paragraph("Algemeen commentaar")
        _spacing(label_para, before=6, after=2)
        for run in label_para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = COLOR_LABEL

        comment_para = doc.add_paragraph(general_comment)
        _spacing(comment_para, before=0, after=4)
        for run in comment_para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_DARK


def _add_page_break(doc: Document) -> None:
    """Insert a hard page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)
    _spacing(para, before=0, after=0)


def _build_doc_base() -> Document:
    """Create a document with narrow margins and Aptos font."""
    doc = Document()
    _set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
    return doc


def _render_team_rubric(
    doc: Document,
    assessment_title: str,
    team_number,
    team_members: list,
    criteria: list,
    scores_map: dict,
    total_score,
    grade,
    general_comment: str | None,
    export_date: str,
) -> None:
    """Render a single team's rubric into an existing document."""
    _add_rubric_header(doc, assessment_title, team_number, team_members, export_date)
    _add_rubric_table(doc, criteria, scores_map)
    _add_score_summary(doc, total_score, grade, general_comment)


# ── Public API ────────────────────────────────────────────────────────────────


def generate_single_team_rubric_docx(data: dict) -> BytesIO:
    """
    Generate a Word document for a single team's rubric.

    Args:
        data: dict with keys:
            assessment_title (str)
            team_number (int|None)
            team_members (list[str])
            criteria (list[dict]): each has id, name, weight, category
            scores_map (dict): criterion_id -> {score, comment}
            total_score (float|None)
            grade (float|None)
            general_comment (str|None)

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = _build_doc_base()
    export_date = date.today().strftime("%d-%m-%Y")

    _render_team_rubric(
        doc,
        assessment_title=data.get("assessment_title") or "",
        team_number=data.get("team_number"),
        team_members=data.get("team_members") or [],
        criteria=data.get("criteria") or [],
        scores_map=data.get("scores_map") or {},
        total_score=data.get("total_score"),
        grade=data.get("grade"),
        general_comment=data.get("general_comment"),
        export_date=export_date,
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_all_teams_rubric_docx(data: dict) -> BytesIO:
    """
    Generate a Word document containing all teams' rubrics (one per page).

    Args:
        data: dict with keys:
            assessment_title (str)
            criteria (list[dict]): each has id, name, weight, category
            teams (list[dict]): each has team_number, team_members, scores_map,
                                total_score, grade, general_comment

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = _build_doc_base()
    export_date = date.today().strftime("%d-%m-%Y")
    assessment_title = data.get("assessment_title") or ""
    criteria = data.get("criteria") or []
    teams = data.get("teams") or []

    for idx, team in enumerate(teams):
        if idx > 0:
            _add_page_break(doc)
        _render_team_rubric(
            doc,
            assessment_title=assessment_title,
            team_number=team.get("team_number"),
            team_members=team.get("team_members") or [],
            criteria=criteria,
            scores_map=team.get("scores_map") or {},
            total_score=team.get("total_score"),
            grade=team.get("grade"),
            general_comment=team.get("general_comment"),
            export_date=export_date,
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
