from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Color palette ──────────────────────────────────────────────────────────────
COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)  # slate-800  – titles, body text
COLOR_ACCENT = RGBColor(0x1E, 0x40, 0xAF)  # blue-800   – section headings
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)  # slate-500  – labels, meta text
COLOR_LABEL = RGBColor(0x47, 0x55, 0x69)  # slate-600  – table labels

_BORDER_COLOR = "CBD5E1"  # slate-300 – table cell borders
_HEADER_FILL = "DBEAFE"  # blue-100  – table header background
_SIG_FILL = "F1F5F9"  # slate-100 – signature header background
_SIG_BORDER = "E2E8F0"  # slate-200 – info / signature borders
_SIGNATURE_LINE_LENGTH = 32  # number of underscores in a signature line

FONT_NAME = "Aptos"  # document-wide typeface

# ── Section ordering & titles ──────────────────────────────────────────────────
SECTION_ORDER = [
    "client",
    "problem",
    "goal",
    "method",
    "planning",
    "tasks",
    "motivation",
    "risks",
]

SECTION_TITLES = {
    "client": "1. Opdrachtgever",
    "problem": "2. Probleemstelling",
    "goal": "3. Doelstelling",
    "method": "4. Methode",
    "planning": "5. Planning",
    "tasks": "6. Taakverdeling",
    "motivation": "7. Motivatie",
    "risks": "8. Risico's",
}


# ── Low-level XML helpers ──────────────────────────────────────────────────────


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Fill a table cell with a background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = _BORDER_COLOR) -> None:
    """Apply light single borders to every edge of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    """Set inner cell padding (values in twentieths of a point / DXA)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_paragraph_bottom_border(para, color: str = "94A3B8") -> None:
    """Add a subtle bottom divider line to a paragraph (slate-400)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _spacing(para, before: float = 0, after: float = 0) -> None:
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def _set_default_font(doc: Document) -> None:
    """Set the document-wide default font to FONT_NAME via the Normal style."""
    # 1. Modify the Normal paragraph style (affects all unstyled text)
    doc.styles["Normal"].font.name = FONT_NAME

    # 2. Patch rPrDefault in docDefaults so the low-level default also uses Aptos.
    #    XML path: w:docDefaults → w:rPrDefault → w:rPr → w:rFonts
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        return
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


# ── Layout helper functions ────────────────────────────────────────────────────


def add_document_header(
    doc: Document,
    project_title: str,
    team_number,
    team_members: list,
    export_date: str,
    client_org: Optional[str] = None,
) -> None:
    """Render the strong top header block: school label → doc type → project title → meta."""
    # Small all-caps school label
    school = doc.add_paragraph("TECHNASIUM MBH")
    school.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(school, before=0, after=2)
    run = school.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED
    run.font.all_caps = True

    # "Projectplan" – large bold document type label
    doc_type = doc.add_paragraph("Projectplan")
    _spacing(doc_type, before=4, after=0)
    run = doc_type.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK

    # Project title in slightly smaller weight
    if project_title:
        proj = doc.add_paragraph(project_title)
        _spacing(proj, before=2, after=6)
        run = proj.runs[0]
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_DARK

    # Metadata line: team · members · date
    parts: list[str] = []
    if team_number is not None:
        parts.append(f"Team {team_number}")
    if team_members:
        parts.append(", ".join(team_members))
    parts.append(f"Datum: {export_date}")
    meta = doc.add_paragraph("   ·   ".join(parts))
    _spacing(meta, before=0, after=2)
    run = meta.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED

    # Optional opdrachtgever hint
    if client_org:
        org_para = doc.add_paragraph(f"Opdrachtgever: {client_org}")
        _spacing(org_para, before=0, after=4)
        run = org_para.runs[0]
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_MUTED
        run.font.italic = True

    # Strong horizontal rule closing the header block
    hr = doc.add_paragraph()
    _spacing(hr, before=6, after=10)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E293B")  # dark slate
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc: Document, text: str) -> None:
    """Bold dark-blue section heading with a subtle bottom divider."""
    para = doc.add_paragraph(text)
    _spacing(para, before=14, after=4)
    run = para.runs[0]
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_ACCENT
    _add_paragraph_bottom_border(para)


def add_info_pairs(doc: Document, pairs: list) -> None:
    """Render label/value pairs as a borderless two-column information block."""
    if not pairs:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in pairs:
        row = table.add_row()
        lc, vc = row.cells[0], row.cells[1]
        lc.text = label
        vc.text = value or ""
        for cell in (lc, vc):
            _set_cell_borders(cell, _SIG_BORDER)
            _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        # Muted bold label
        if lc.paragraphs[0].runs:
            r = lc.paragraphs[0].runs[0]
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = COLOR_LABEL
        # Dark value
        if vc.paragraphs[0].runs:
            r = vc.paragraphs[0].runs[0]
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK
    spacer = doc.add_paragraph()
    _spacing(spacer, before=0, after=4)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Readable body paragraph with comfortable line spacing."""
    para = doc.add_paragraph(text)
    _spacing(para, before=0, after=8)
    para.paragraph_format.line_spacing = Pt(14)
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK


def add_bullet_list(doc: Document, items: list) -> None:
    """Consistent bullet list items."""
    for item in items:
        if not item.strip():
            continue
        para = doc.add_paragraph(item.strip(), style="List Bullet")
        _spacing(para, before=1, after=1)
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_DARK


def add_styled_table(doc: Document, header_labels: list, rows: list) -> None:
    """Neat table: tinted header row, light borders, comfortable cell padding."""
    if not rows:
        return
    col_count = max(len(header_labels), 1)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    if header_labels:
        hdr_row = table.add_row()
        for i, label in enumerate(header_labels):
            if i >= col_count:
                break
            cell = hdr_row.cells[i]
            cell.text = label
            _set_cell_shading(cell, _HEADER_FILL)
            _set_cell_borders(cell, _BORDER_COLOR)
            _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.name = FONT_NAME
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_DARK

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            if i >= col_count:
                break
            cell = row.cells[i]
            cell.text = str(value) if value else ""
            _set_cell_borders(cell, _BORDER_COLOR)
            _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_DARK

    spacer = doc.add_paragraph()
    _spacing(spacer, before=0, after=6)


def add_signature_section(doc: Document) -> None:
    """Clean two-column signature block: Opdrachtgever | Projectteam."""
    intro = doc.add_paragraph(
        "Hierbij verklaren ondergetekenden akkoord te gaan met de inhoud van "
        "bovenstaand projectplan."
    )
    _spacing(intro, before=0, after=12)
    for run in intro.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK

    sig_table = doc.add_table(rows=0, cols=2)
    sig_table.style = "Table Grid"

    # Column header row
    hdr_row = sig_table.add_row()
    for i, label in enumerate(("Opdrachtgever", "Projectteam")):
        cell = hdr_row.cells[i]
        cell.text = label
        _set_cell_shading(cell, _SIG_FILL)
        _set_cell_borders(cell, _SIG_BORDER)
        _set_cell_margins(cell, top=80, bottom=80, left=160, right=160)
        if cell.paragraphs[0].runs:
            r = cell.paragraphs[0].runs[0]
            r.font.name = FONT_NAME
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = COLOR_DARK

    # Naam / Handtekening / Datum rows
    for field_label in ("Naam:", "Handtekening:", "Datum:"):
        data_row = sig_table.add_row()
        for cell in data_row.cells:
            para = cell.paragraphs[0]
            para.clear()
            r = para.add_run(field_label)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = COLOR_LABEL
            # Writing space
            space_para = cell.add_paragraph()
            _spacing(space_para, before=14, after=4)
            sig_run = space_para.add_run("_" * _SIGNATURE_LINE_LENGTH)
            sig_run.font.name = FONT_NAME
            sig_run.font.size = Pt(10)
            sig_run.font.color.rgb = COLOR_MUTED
            _set_cell_borders(cell, _SIG_BORDER)
            _set_cell_margins(cell, top=100, bottom=120, left=160, right=160)


# ── Text parsing helpers ───────────────────────────────────────────────────────


def _parse_lines(text: str) -> list:
    """Return non-empty stripped lines from a block of text."""
    return [ln.strip() for ln in text.splitlines() if ln.strip()]


# ── Main entry point ───────────────────────────────────────────────────────────


def generate_projectplan_docx(team_data: dict) -> BytesIO:
    """
    Generate a professionally styled Word document for a projectplan.

    Args:
        team_data: dict with title, team_number, team_members, sections

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Set document-wide default font
    _set_default_font(doc)

    # Slightly tighter margins for a cleaner print layout
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_text = team_data.get("title") or "Projectplan"
    team_number = team_data.get("team_number")
    team_members = team_data.get("team_members") or []
    export_date = date.today().strftime("%d-%m-%Y")
    sections_map = {s["key"]: s for s in team_data.get("sections", [])}

    # Extract client org name for the header hint
    client_org: Optional[str] = None
    client_section = sections_map.get("client")
    if client_section:
        client_org = (client_section.get("client") or {}).get("organisation") or None

    # ── 1. Header block ────────────────────────────────────────────────────────
    add_document_header(
        doc, title_text, team_number, team_members, export_date, client_org
    )

    # ── 2. Content sections ────────────────────────────────────────────────────
    for key in SECTION_ORDER:
        section = sections_map.get(key)
        if not section:
            continue

        # ── Opdrachtgever ──────────────────────────────────────────────────────
        if key == "client":
            client = section.get("client") or {}
            info_pairs = [
                ("Organisatie", client.get("organisation")),
                ("Contactpersoon", client.get("contact")),
                ("Email", client.get("email")),
                ("Telefoon", client.get("phone")),
            ]
            non_empty = [(lbl, val) for lbl, val in info_pairs if val]
            description = (client.get("description") or "").strip()
            if not non_empty and not description:
                continue
            add_section_heading(doc, SECTION_TITLES[key])
            if non_empty:
                add_info_pairs(doc, non_empty)
            if description:
                add_body_paragraph(doc, description)

        # ── Planning / Risico's → styled table ────────────────────────────────
        elif key in ("planning", "risks"):
            text = (section.get("text") or "").strip()
            if not text:
                continue
            add_section_heading(doc, SECTION_TITLES[key])
            lines = _parse_lines(text)
            if lines:
                hdr = ["Risico"] if key == "risks" else ["Periode / Activiteit"]
                add_styled_table(doc, hdr, [[ln] for ln in lines])
            else:
                add_body_paragraph(doc, text)

        # ── Methode → bold sub-headings + bullets ─────────────────────────────
        elif key == "method":
            text = (section.get("text") or "").strip()
            if not text:
                continue
            add_section_heading(doc, SECTION_TITLES[key])
            # Double-newline separates named sub-parts; single newline = bullets
            blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
            if len(blocks) <= 1:
                lines = _parse_lines(text)
                if len(lines) > 1:
                    add_bullet_list(doc, lines)
                else:
                    add_body_paragraph(doc, text)
            else:
                for block in blocks:
                    sub_lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
                    if not sub_lines:
                        continue
                    if len(sub_lines) == 1:
                        add_body_paragraph(doc, sub_lines[0])
                    else:
                        # First line of a block → bold sub-heading
                        subhead = doc.add_paragraph(sub_lines[0])
                        _spacing(subhead, before=8, after=2)
                        for run in subhead.runs:
                            run.font.name = FONT_NAME
                            run.font.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = COLOR_DARK
                        add_bullet_list(doc, sub_lines[1:])

        # ── Probleemstelling / Doelstelling / Motivatie → body + bullets ──────
        elif key in ("problem", "goal", "motivation"):
            text = (section.get("text") or "").strip()
            if not text:
                continue
            add_section_heading(doc, SECTION_TITLES[key])
            lines = _parse_lines(text)
            if len(lines) > 1:
                add_body_paragraph(doc, lines[0])
                add_bullet_list(doc, lines[1:])
            else:
                add_body_paragraph(doc, text)

        # ── All other sections (Taakverdeling, …) → body or bullets ─────────
        else:
            text = (section.get("text") or "").strip()
            if not text:
                continue
            add_section_heading(doc, SECTION_TITLES[key])
            lines = _parse_lines(text)
            if len(lines) > 1:
                add_body_paragraph(doc, lines[0])
                add_bullet_list(doc, lines[1:])
            else:
                add_body_paragraph(doc, text)

    # ── 3. Akkoordverklaring ───────────────────────────────────────────────────
    add_section_heading(doc, "Akkoordverklaring")
    add_signature_section(doc)

    # ── Save ──────────────────────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
